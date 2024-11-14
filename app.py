import streamlit as st
from environment import detect_system_resources
import os
import time

# Now import other modules that depend on installed packages
from preprocessing import preprocess_text
from visualization import create_multi_spider_chart, chart_to_image
from synergy import generate_synergy_report
import tempfile
from docx import Document
from docx.shared import Inches
from io import BytesIO

# Import all dimension modules (e.g., Affective, Animate, etc.)
from dimensions import (
    Affective, Animate, Commonality, Cognitive, Directness, Dynamic, Formality, Generic,
    Inanimate, Individual, Informality, Intentionality, LongTerm, Negative, Novelty,
    Objectivity, Place, Politeness, Positive, Qualitative, Quantitative, ShortTerm,
    Social, Specific, Static, Time
)
from llm_loader import load_llm_model, get_embedding

# Create a mapping of dimension names to their respective class instances
DIMENSIONS = {
    'Affective': Affective(),
    'Animate': Animate(),
    'Commonality': Commonality(),
    'Cognitive': Cognitive(),
    'Directness': Directness(),
    'Dynamic': Dynamic(),
    'Formality': Formality(),
    'Generic': Generic(),
    'Inanimate': Inanimate(),
    'Individual': Individual(),
    'Informality': Informality(),
    'Intentionality': Intentionality(),
    'LongTerm': LongTerm(),
    'Negative': Negative(),
    'Novelty': Novelty(),
    'Objectivity': Objectivity(),
    'Place': Place(),
    'Politeness': Politeness(),
    'Positive': Positive(),
    'Qualitative': Qualitative(),
    'Quantitative': Quantitative(),
    'ShortTerm': ShortTerm(),
    'Social': Social(),
    'Specific': Specific(),
    'Static': Static(),
    'Time': Time(),
}

def split_text_into_chunks(text, chunk_size=1000000):
    """
    Splits the text into smaller chunks to handle processing of large texts.

    :param text: The input text to split.
    :param chunk_size: Maximum size of each chunk.
    :return: A list of text chunks.
    """
    return [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

def calculate_dimension_scores(text, method='standard'):
    """
    Calculate scores for all dimensions for the input text.

    :param text: The input text to analyze.
    :param method: 'standard' or 'advanced'.
    :return: A dictionary of dimension scores with error flags.
    """
    dimension_scores = {}
    total_dimensions = len(DIMENSIONS)
    progress_bar = st.progress(0)
    eta_placeholder = st.empty()
    start_time = time.time()

    for idx, (dimension_name, dimension_class) in enumerate(DIMENSIONS.items()):
        try:
            if method == 'advanced':
                # Load the relevant LLM for the dimension
                model = load_llm_model(dimension_name)
                score_dict = dimension_class.calculate_score(text, method='advanced', model=model)
            else:
                score_dict = dimension_class.calculate_score(text, method='standard')

            # Ensure score_dict has the dimension_name as key with 'score' and 'error'
            if dimension_name in score_dict:
                dimension_scores[dimension_name] = {
                    'score': score_dict[dimension_name]['score'],
                    'error': score_dict[dimension_name]['error']
                }
            else:
                # Handle unexpected dictionary structure
                dimension_scores[dimension_name] = {
                    'score': 0.0,
                    'error': True,
                    'error_message': 'Unexpected score structure.'
                }

            # **Add error-checking code here**
            details = dimension_scores[dimension_name]
            if details.get('error', False):
                error_message = details.get('error_message', 'No error message provided.')
                st.error(f"Error in {dimension_name}: {error_message}")

        except Exception as e:
            st.warning(f"Error calculating score for {dimension_name}: {e}. Reverting to standard method.")
            try:
                # Fallback to standard method
                score_dict = dimension_class.calculate_score(text, method='standard')
                if dimension_name in score_dict:
                    dimension_scores[dimension_name] = {
                        'score': score_dict[dimension_name]['score'],
                        'error': score_dict[dimension_name]['error']
                    }
                else:
                    dimension_scores[dimension_name] = {
                        'score': 0.0,
                        'error': True,
                        'error_message': 'Unexpected score structure in standard method.'
                    }
                st.info(f"Standard method used for {dimension_name} due to the following error in advanced method: {e}")

                # **Add error-checking code here for standard method fallback**
                details = dimension_scores[dimension_name]
                if details.get('error', False):
                    error_message = details.get('error_message', 'No error message provided.')
                    st.error(f"Error in {dimension_name}: {error_message}")

            except Exception as standard_error:
                st.error(f"Error calculating standard score for {dimension_name}: {standard_error}")
                dimension_scores[dimension_name] = {
                    'score': 0.0,
                    'error': True,
                    'error_message': str(standard_error)
                }

                # **Add error-checking code here for double exception**
                details = dimension_scores[dimension_name]
                if details.get('error', False):
                    error_message = details.get('error_message', 'No error message provided.')
                    st.error(f"Error in {dimension_name}: {error_message}")

        # Update progress bar and ETA
        progress = (idx + 1) / total_dimensions
        progress_bar.progress(progress)
        elapsed_time = time.time() - start_time
        average_time_per_dimension = elapsed_time / (idx + 1)
        remaining_dimensions = total_dimensions - (idx + 1)
        eta = average_time_per_dimension * remaining_dimensions
        eta_placeholder.text(f"Estimated time remaining: {int(eta)} seconds")

    # Ensure all expected dimensions are present
    expected_dimensions = list(DIMENSIONS.keys())
    for dim in expected_dimensions:
        if dim not in dimension_scores:
            dimension_scores[dim] = {
                'score': 0.0,
                'error': False
            }
            st.warning(f"Score for dimension '{dim}' was missing and has been set to 0.0.")

    eta_placeholder.text("Analysis complete!")
    return dimension_scores

def generate_overall_interpretation(dimension_scores):
    """
    Generates an overall interpretation of the results based on dimension scores.

    :param dimension_scores: A dictionary of dimension scores.
    :return: A summary string.
    """
    sorted_scores = sorted(dimension_scores.items(), key=lambda x: x[1]['score'], reverse=True)
    highest_dimensions = [dim for dim, details in sorted_scores[:3] if not details['error']]
    lowest_dimensions = [dim for dim, details in sorted_scores[-3:] if not details['error']]

    interpretation = "Overall Interpretation:\n"
    if highest_dimensions:
        interpretation += "The text shows strong indications in the following dimensions:\n"
        for dim in highest_dimensions:
            interpretation += f"- {dim} (Score: {dimension_scores[dim]['score']:.2f})\n"
    else:
        interpretation += "No dimensions show particularly strong indications.\n"

    if lowest_dimensions:
        interpretation += "\nThe text shows weaker indications in the following dimensions:\n"
        for dim in lowest_dimensions:
            interpretation += f"- {dim} (Score: {dimension_scores[dim]['score']:.2f})\n"
    else:
        interpretation += "\nNo dimensions show particularly weak indications.\n"

    return interpretation

def generate_final_report(dimension_scores, input_name="Input", method='standard'):
    """
    Generate a final analysis report including dimension scores, synergy analysis, and overall interpretation.

    :param dimension_scores: A dictionary of dimension scores with error flags.
    :param input_name: Name of the input source.
    :param method: The method used ('standard' or 'advanced').
    :return: A string representing the final report.
    """
    report = f"{input_name} Analysis Report:\n"
    report += "---------------------\n"

    # Add dimension scores to the report
    for dimension, details in dimension_scores.items():
        if 'score' in details:
            report += f"{dimension}: {details['score']:.2f}\n"
        else:
            report += f"{dimension}: N/A (Error in analysis)\n"

    # Generate synergy report
    synergy_report = generate_synergy_report(dimension_scores)
    report += f"\n{synergy_report}\n\n"

    # Generate overall interpretation
    interpretation = generate_overall_interpretation(dimension_scores)
    report += f"{interpretation}\n\n"

    # Add method information
    report += f"**Analysis Method Used:** {'Advanced' if method == 'advanced' else 'Standard'}\n"
    report += "Interpretations are based on the dimension scores calculated using the selected method.\n"

    # Add Zero Scores Interpretation
    report += "\n## Zero Scores Interpretation:\n"
    zero_dims = [dim for dim, details in dimension_scores.items() if details.get('score', 0.0) == 0.0]
    if zero_dims:
        report += "The following dimensions have a score of 0.00:\n"
        for dim in zero_dims:
            if dimension_scores[dim].get('error', False):
                report += f"- **{dim}**: Score is 0.00 due to an error during analysis.\n"
            else:
                report += f"- **{dim}**: Score is 0.00 indicating the text lacks characteristics related to this dimension.\n"
    else:
        report += "No dimensions have a score of 0.00.\n"

    return report

def create_docx_report(report_text, input_name, individual_chart, combined_chart):
    """
    Creates a DOCX report containing the analysis text and charts.

    :param report_text: The analysis report text.
    :param input_name: Name of the input source.
    :param individual_chart: Bytes object of the individual radar chart image.
    :param combined_chart: Bytes object of the combined radar chart image.
    :return: BytesIO object of the DOCX file.
    """
    doc = Document()
    doc.add_heading(f'Analysis Report for {input_name}', level=1)
    doc.add_paragraph(report_text)

    doc.add_heading('Radar Chart for This Input', level=2)
    doc.add_picture(BytesIO(individual_chart), width=Inches(6))

    doc.add_heading('Combined Radar Chart for All Inputs', level=2)
    doc.add_picture(BytesIO(combined_chart), width=Inches(6))

    # Save the document to a BytesIO object
    f = BytesIO()
    doc.save(f)
    f.seek(0)
    return f

def estimate_required_memory(text_length, method='advanced'):
    """
    Estimates the required memory for processing based on text length and method.

    :param text_length: Total length of input text.
    :param method: The analysis method ('standard' or 'advanced').
    :return: Estimated memory in GB.
    """
    base_memory = 1  # Base memory usage in GB
    if method == 'advanced':
        # LLMs require more memory, estimate based on text length
        memory_per_char = 0.00002  # Approximate memory per character
    else:
        # Standard methods are less memory-intensive
        memory_per_char = 0.000005
    estimated_memory = base_memory + (text_length * memory_per_char)
    return estimated_memory

def main():
    st.title("Multidimensional Text Analysis")

    # Initialize session state for analysis results if not already present
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = {}

    # Detect system resources
    resources = detect_system_resources()
    st.sidebar.write("**System Resources Detected:**")
    st.sidebar.write(f"RAM: {resources['ram']:.2f} GB")
    st.sidebar.write(f"CPU Cores: {resources['cpu']}")
    st.sidebar.write(f"GPU Available: {'Yes' if resources['gpu'] else 'No'}")

    st.write("""
    Welcome to the Multidimensional Text Analysis Tool! Upload one or more text files, PDFs, DOCX files, or enter URLs to analyze the content across multiple dimensions. The app will handle all processing behind the scenes.
    """)

    # Allow multiple file uploads
    uploaded_files = st.file_uploader(
        "Choose file(s) to analyze",
        type=["txt", "csv", "docx", "pdf"],
        accept_multiple_files=True
    )
    url_inputs = st.text_area(
        "Or enter URLs to analyze (one per line)",
        placeholder="https://example.com/article1\nhttps://example.com/article2"
    )

    # Collect all input sources
    input_sources = []

    if uploaded_files:
        input_sources.extend(uploaded_files)

    if url_inputs.strip():
        urls = [url.strip() for url in url_inputs.strip().split('\n') if url.strip()]
        input_sources.extend(urls)

    if input_sources:
        # Preprocess inputs to get text lengths and complexities
        input_texts = {}
        total_text_length = 0
        for source in input_sources:
            # Get input name
            input_name = source if isinstance(source, str) else source.name
            try:
                text = preprocess_text(source)
                input_texts[input_name] = text
                total_text_length += len(text)
            except ValueError as e:
                st.error(f"Error processing {input_name}: {e}")
                continue

        st.sidebar.write(f"**Total Input Text Length:** {total_text_length} characters")

        # Estimate required resources
        estimated_memory = estimate_required_memory(total_text_length, method='advanced')

        st.sidebar.write(f"**Estimated Memory Required for Advanced Method:** {estimated_memory:.2f} GB")

        # Compare with available resources
        if estimated_memory > resources['ram']:
            st.sidebar.warning("The advanced method may exceed your system's memory capacity.")
            recommended_method = 'standard'
        else:
            st.sidebar.success("Your system resources are sufficient for the advanced method.")
            recommended_method = 'advanced'

        # Allow user to select method
        method = st.sidebar.radio(
            "Select Analysis Method:",
            ('Standard (Faster, Less Accurate)', 'Advanced (Slower, More Accurate)'),
            index=0 if recommended_method == 'standard' else 1
        )

        method_value = 'advanced' if method == 'Advanced (Slower, More Accurate)' else 'standard'

        if st.button("Start Analysis"):
            all_dimension_scores = {}
            individual_charts = {}
            with st.spinner("Processing... Please wait."):
                for input_name, text in input_texts.items():
                    # Check if text is empty
                    if not text.strip():
                        st.warning(f"The input text '{input_name}' is empty after preprocessing.")
                        continue

                    # Split text into chunks if necessary
                    if len(text) > 1000000:
                        text_chunks = split_text_into_chunks(text)
                        st.info(f"Splitting '{input_name}' into {len(text_chunks)} chunks for processing.")
                    else:
                        text_chunks = [text]

                    # Initialize dimension scores with error flags
                    dimension_scores = {}
                    for idx, chunk in enumerate(text_chunks):
                        st.write(f"Processing chunk {idx + 1}/{len(text_chunks)} for '{input_name}'...")
                        try:
                            chunk_scores = calculate_dimension_scores(chunk, method=method_value)
                        except Exception as e:
                            st.warning(f"Advanced method failed for chunk {idx + 1}: {e}")
                            chunk_scores = calculate_dimension_scores(chunk, method='standard')
                            st.info(f"Falling back to Standard method for chunk {idx + 1}")

                        # Aggregate scores with error tracking
                        for dimension, details in chunk_scores.items():
                            if dimension in dimension_scores:
                                # Aggregate scores by summing
                                dimension_scores[dimension]['score'] += details.get('score', 0.0)
                                # If any chunk has an error, mark the dimension as having an error
                                dimension_scores[dimension]['error'] = dimension_scores[dimension]['error'] or details.get('error', False)
                            else:
                                dimension_scores[dimension] = {
                                    'score': details.get('score', 0.0),
                                    'error': details.get('error', False)
                                }

                    # Average the scores over the chunks
                    for dimension in dimension_scores:
                        dimension_scores[dimension]['score'] /= len(text_chunks)

                    all_dimension_scores[input_name] = dimension_scores

                    # Generate individual radar chart with only scores
                    scores_only = {dim: details['score'] for dim, details in dimension_scores.items()}
                    individual_chart = create_multi_spider_chart({input_name: scores_only}, show=False)
                    individual_charts[input_name] = individual_chart

                    # Generate synergy report for the input
                    synergy_report = generate_synergy_report(dimension_scores)

                    # Generate final report
                    final_report = generate_final_report(dimension_scores, input_name, method=method_value)

                    # Store in session state
                    st.session_state.analysis_results[input_name] = {
                        'scores': dimension_scores,
                        'report': final_report,
                        'synergy': synergy_report,
                        'individual_chart': individual_chart
                    }

            if all_dimension_scores:
                # Generate combined radar chart with only scores
                all_scores_only = {input_name: {dim: details['score'] for dim, details in scores.items()} 
                                   for input_name, scores in all_dimension_scores.items()}
                combined_chart = create_multi_spider_chart(all_scores_only, show=False)

                # Store combined chart in session state
                st.session_state.combined_chart = combined_chart

                # Display reports and charts
                for input_name, scores in all_dimension_scores.items():
                    report = st.session_state.analysis_results[input_name]['report']
                    st.subheader(f"Analysis Report for {input_name}")
                    st.text_area(f"Report for {input_name}", report, height=300)

                    # Prepare charts for the report
                    individual_chart_img = chart_to_image(individual_charts[input_name])
                    combined_chart_img = chart_to_image(combined_chart)

                    # Create DOCX report
                    docx_report = create_docx_report(report, input_name, individual_chart_img, combined_chart_img)

                    # Download report as a DOCX file
                    st.download_button(
                        label=f"Download {input_name} Report",
                        data=docx_report,
                        file_name=f"{input_name}_analysis_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                # Display combined radar chart
                st.subheader("Combined Radar Chart for All Inputs")
                st.pyplot(combined_chart)

                # Download combined chart as an image
                combined_chart_img = chart_to_image(combined_chart)
                st.download_button(
                    label="Download Combined Radar Chart",
                    data=combined_chart_img,
                    file_name="combined_dimension_radar_chart.png",
                    mime="image/png"
                )

                # Generate and display general report for all files
                general_report = "## General Analysis Report for All Inputs:\n"
                for input_name, data in st.session_state.analysis_results.items():
                    general_report += f"### {input_name}:\n"
                    general_report += f"{data['report']}\n\n"

                st.subheader("General Analysis Report for All Inputs")
                st.text_area("General Report", general_report, height=600)

                # Create a comprehensive DOCX report for all inputs
                comprehensive_docx = create_docx_report(
                    general_report,
                    "All Inputs",
                    combined_chart_img,  # You can customize this as needed
                    combined_chart_img
                )

                st.download_button(
                    label="Download Comprehensive Report for All Inputs",
                    data=comprehensive_docx,
                    file_name="comprehensive_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.success("Analysis complete!")
            else:
                st.error("No valid inputs were provided for analysis.")

if __name__ == "__main__":
    main()

